import whisper
import os
from tqdm import tqdm

import ollama
from ollama import chat
from ollama import ChatResponse

from docx import Document

import librosa

import warnings


"""
This project is used to transcribe audio files to text using the OpenAI Whisper model.
It uses a virtual environment with the requirements installed there.
It is intended to be used to transcribe audio files and then convert the .txt files to a .docx file.
From here, a LLM will be used to edit the .docx file to make it more readable.
"""

def fxn():
    warnings.warn("deprecated", DeprecationWarning)


def transcribe_audio(model, recording_path, length_of_audio):
    """_summary_

    Args:
        model: The Whisper model to use

    Returns:
        str: The transcribed audio text as a string 
    """    
    print(f"Transcribing your audio now from {recording_path}...")
    result = model.transcribe(recording_path)
    
    # Filter segments to within audio duration if the segment start time is after the audio is over
    segments = result["segments"]
    trimmed_segments = [seg for seg in segments if seg["start"] <= length_of_audio]

    # Optional: combine back into a single string
    final_transcript = " ".join(seg["text"] for seg in trimmed_segments)

    return final_transcript

def write_text_to_file(item, result, recording_text_files):
    """_summary_

    Args:
        text: The text to write to the file
        result: The result of the transcribed audio
        recording_text_files: The path to the file to write to
        
    Returns:
        None
        Saves a .txt file of the transcribed audio
    """    
    print(f"Writing {item} to a .txt file...")
    name = item.split(".")[0]
    #Need to add encoding incase the recording has special characters
    with open(f"{recording_text_files}/{name}.txt", "w", encoding="utf-8") as f:
        f.write(result)

def correct_grammar_and_edit(item, result, recording_grammar_file_path, recording_ai_file_path):
    print(f"Correcting the grammar in the {item} text")
    # Chat with the Llama model to correct the grammar
    response: ChatResponse = chat(model='llama3.2', messages=[
    {
        'role': 'user',
        'content': f'Please correct the grammar, spelling and punctuation in the {result} text. Word alteration should be minimal at most. Return only the corrected text, no formatting or comments.',
    },
    ])

    #print(response.message.content)

    # Save the corrected text to a .docx file
    document = Document()
    name = item.split(".")[0]

    document.add_heading(f'{name}', level=1)
    document.add_paragraph(response.message.content)

    document.save(f'{recording_grammar_file_path}/{name}.docx')
    print(f"Your text has been corrected and saved to the {name}.docx file in the {recording_grammar_file_path} folder.")

    # Chat with the Llama model to edit the document
    ai_edits(item, response, recording_ai_file_path)
    
def ai_edits(item, response, recording_ai_file_path):
    print(f"Editing the {item} text with the Llama ai model")
    new_response: ChatResponse = chat(model='llama3.2', messages=[
    {
        'role': 'user',
        'content': f'You are a clarity and editing assistant. Your job is to correct grammar, punctuation, and flow while keeping the original tone and word choices as much as possible while making the writing sound more professional. Only return the corrected version of this input. Do not add commentary, summaries, introductions, or explanations. Input:{response}',
    },
    ])

    # Save the corrected text to a .docx file
    document = Document()
    name = item.split(".")[0]

    document.add_heading(f'{name}', level=1)
    document.add_paragraph(new_response.message.content)

    document.save(f'{recording_ai_file_path}/{name}.docx')
    print(f"Your text has been edited and saved to the {name}.docx file in the {recording_ai_file_path} folder.")


if __name__ == "__main__":
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        fxn()
    
        print("Welcome to the Whisper program! The program will transcribe your audio files to text and then correct the grammar and edit the text. It is currently loading the Whisper model and the Llama model. Please wait a moment.")
        model = whisper.load_model("base") #load whisper model: tiny, base, small, medium, large, or turbo, see https://github.com/openai/whisper
        
        recordings = "./Recordings" #path to audio file
        recording_text_files = "./Recording_text_files" #path to save text files
        recording_grammar_file_path = "./Recording_docx_files_grammar_edits" #path to save corrected .docx files
        recording_ai_file_path = "./Recording_docx_files_AI_edits" #path to save AI corrected .docx files
        
        recording_titles = os.listdir(recordings) #list of audio files
        
        # NOTE: The Llama model is not available in the virtual environment and must be downloaded. 
        # Mistral (or whatever other model you choose) must be downloaded as well.
        # Please see: https://github.com/ollama/ollama-python/blob/main/README.md for more information.
        
        # in Windows I had to run the following commands in the Powershell to run/install the Llama model:
        # ollama serve
        # ollama pull mistral
        # ollama pull llama3.2
        
        # Load the Llama model
        ollama_model = 'mistral'

        try:
            ollama.chat(ollama_model)
            print('Ollama model is available and running.')
        except ollama.ResponseError as e:
            print('Error:', e.error)
            if e.status_code == 404:
                ollama.pull(ollama_model)


        for item in tqdm(recording_titles):
            story_path = os.path.join(recordings, item)
            
            # Get the length of the audio file to be able to truncate and avoid ollama halucinations
            length_of_audio = librosa.get_duration(filename=story_path)
            print(f"\nYour audio is {length_of_audio} seconds long.")

            # STEP 1: transcribe the audio
            result = transcribe_audio(model, story_path, length_of_audio)
            print("Congratulations! Your audio has been transcribed.\n")
            
            # STEP 2: write the text to a .txt file
            write_text_to_file(item, result, recording_text_files)
            print("Your text has been written to a .txt file.\n")
            
            # STEP 3: correct the grammar
            # NOTE: The correct_grammar function will also call the ai_edits function to edit the document separately
            correct_grammar_and_edit(item, result, recording_grammar_file_path, recording_ai_file_path)
            print("Your text has been corrected and edited. Thanks for using the program :)\n")